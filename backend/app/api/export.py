"""
Export API — Generate PDF and DOCX reports from contract analysis.
"""
import os
import logging
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session

from app.database import get_db
from app import models
from app.utils import get_current_user

router = APIRouter()
logger = logging.getLogger(__name__)

EXPORT_DIR = os.getenv("EXPORT_DIR", "exports")


def _generate_pdf_report(contract: models.Contract, clauses: list, output_path: str):
    """Generate PDF analysis report using ReportLab."""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=0.75*inch, bottomMargin=0.75*inch)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=22, spaceAfter=6, textColor=colors.HexColor("#0F172A"))
        story.append(Paragraph("Contract Analysis Report", title_style))
        story.append(Paragraph(f"<b>{contract.title}</b>", styles["Normal"]))
        story.append(Spacer(1, 12))

        # Risk Score
        score = contract.risk_score or 50
        score_color = "#10B981" if score >= 75 else "#F59E0B" if score >= 50 else "#EF4444"
        story.append(Paragraph(f"<b>Contract Health Score: <font color='{score_color}'>{score}/100</font></b>", styles["Heading2"]))
        story.append(Spacer(1, 8))

        # Stats table
        high = sum(1 for c in clauses if c.risk_level == "high")
        medium = sum(1 for c in clauses if c.risk_level == "medium")
        low = sum(1 for c in clauses if c.risk_level == "low")

        data = [
            ["Total Clauses", "High Risk", "Medium Risk", "Low Risk"],
            [str(len(clauses)), str(high), str(medium), str(low)],
        ]
        t = Table(data, colWidths=[1.5*inch]*4)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F8FAFC"), colors.white]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ("ROUNDEDCORNERS", [5, 5, 5, 5]),
        ]))
        story.append(t)
        story.append(Spacer(1, 16))

        # Clauses
        story.append(Paragraph("Detected Clauses", styles["Heading2"]))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E2E8F0")))
        story.append(Spacer(1, 8))

        for clause in clauses:
            risk_color = "#EF4444" if clause.risk_level == "high" else "#F59E0B" if clause.risk_level == "medium" else "#10B981"
            story.append(Paragraph(
                f"<b>{clause.clause_type}</b> — <font color='{risk_color}'>{clause.risk_level.upper()} RISK</font>",
                styles["Heading3"]
            ))
            if clause.explanation:
                story.append(Paragraph(f"<b>Explanation:</b> {clause.explanation}", styles["Normal"]))
            if clause.optimized_text:
                story.append(Paragraph(f"<b>Optimized:</b> {clause.optimized_text[:300]}", styles["Normal"]))
            if clause.negotiation_tip:
                story.append(Paragraph(f"<b>Tip:</b> {clause.negotiation_tip}", styles["Normal"]))
            story.append(Spacer(1, 10))

        doc.build(story)
        return True
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return False


def _generate_docx_report(contract: models.Contract, clauses: list, output_path: str):
    """Generate DOCX analysis report."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading("Contract Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading(contract.title, 1)

        score = contract.risk_score or 50
        p = doc.add_paragraph()
        p.add_run(f"Contract Health Score: {score}/100").bold = True

        doc.add_paragraph(f"Total Clauses: {len(clauses)}")
        high = sum(1 for c in clauses if c.risk_level == "high")
        medium = sum(1 for c in clauses if c.risk_level == "medium")
        low = sum(1 for c in clauses if c.risk_level == "low")
        doc.add_paragraph(f"High Risk: {high} | Medium Risk: {medium} | Low Risk: {low}")

        doc.add_page_break()
        doc.add_heading("Clause Analysis", 1)

        for clause in clauses:
            doc.add_heading(f"{clause.clause_type} — {clause.risk_level.upper()} RISK", 2)
            if clause.explanation:
                p = doc.add_paragraph()
                p.add_run("Explanation: ").bold = True
                p.add_run(clause.explanation)
            if clause.optimized_text:
                p = doc.add_paragraph()
                p.add_run("Optimized Version: ").bold = True
                p.add_run(clause.optimized_text[:500])
            if clause.negotiation_tip:
                p = doc.add_paragraph()
                p.add_run("Negotiation Tip: ").bold = True
                p.add_run(clause.negotiation_tip)
            doc.add_paragraph("—" * 40)

        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return False


@router.get("/pdf/{contract_id}")
async def export_pdf(
    contract_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Export contract analysis as PDF."""
    contract = db.query(models.Contract).filter(
        models.Contract.id == contract_id,
        models.Contract.user_id == current_user.id,
    ).first()

    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    clauses = db.query(models.Clause).filter(models.Clause.contract_id == contract_id).all()

    os.makedirs(EXPORT_DIR, exist_ok=True)
    output_path = os.path.join(EXPORT_DIR, f"analysis_{contract_id}.pdf")

    success = _generate_pdf_report(contract, clauses, output_path)
    if not success or not os.path.exists(output_path):
        raise HTTPException(status_code=500, detail="Failed to generate PDF")

    return FileResponse(
        output_path,
        media_type="application/pdf",
        filename=f"analysis_{contract.title[:30]}.pdf",
    )


@router.get("/docx/{contract_id}")
async def export_docx(
    contract_id: int,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Export contract analysis as DOCX."""
    contract = db.query(models.Contract).filter(
        models.Contract.id == contract_id,
        models.Contract.user_id == current_user.id,
    ).first()

    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    clauses = db.query(models.Clause).filter(models.Clause.contract_id == contract_id).all()

    os.makedirs(EXPORT_DIR, exist_ok=True)
    output_path = os.path.join(EXPORT_DIR, f"analysis_{contract_id}.docx")

    success = _generate_docx_report(contract, clauses, output_path)
    if not success or not os.path.exists(output_path):
        raise HTTPException(status_code=500, detail="Failed to generate DOCX")

    return FileResponse(
        output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"analysis_{contract.title[:30]}.docx",
    )
